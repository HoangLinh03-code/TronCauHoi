# -*- coding: utf-8 -*-
"""
Debug script: Xem chi tiết XML formatting của từng run trong DOCX.
Mục đích: Tìm hiểu tại sao phát hiện đáp án đúng (đỏ/gạch chân) bị sai.
"""

import sys
import os
import re
from lxml import etree
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def dump_paragraph_formatting(filepath, output_path=None):
    """
    In chi tiết formatting từng run trong file DOCX.
    Hỗ trợ tùy chọn ghi kết quả ra file text thông qua tham số output_path.
    """
    # Mở file để ghi nếu có yêu cầu (mode 'w' để ghi đè)
    f_out = open(output_path, 'w', encoding='utf-8') if output_path else None
    
    # Hàm helper nội bộ: Vừa in ra console, vừa ghi vào file
    def log(msg=""):
        print(msg)
        if f_out:
            f_out.write(str(msg) + "\n")

    try:
        doc = Document(filepath)
        
        log(f"\n{'='*70}")
        log(f"  FILE: {os.path.basename(filepath)}")
        log(f"  Tổng paragraphs: {len(doc.paragraphs)}")
        log(f"{'='*70}")
        
        # Pattern nhận diện đáp án và câu hỏi
        opt_pattern = re.compile(r'^([A-D])\s*[.)]\s*', re.IGNORECASE)
        q_pattern = re.compile(r'(?:Câu|Question)\s*(\d+)', re.IGNORECASE)
        
        current_q = None
        
        for pi, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Kiểm tra xem có phải câu hỏi mới không
            qm = q_pattern.match(text)
            if qm:
                current_q = int(qm.group(1))
                log(f"\n{'─'*70}")
                log(f"  📌 Câu {current_q}")
            
            # Kiểm tra đáp án
            is_option = opt_pattern.match(text)
            
            if is_option or qm:
                prefix = f"  [Đáp án {is_option.group(1)}]" if is_option else f"  [Stem]"
                log(f"\n{prefix} Paragraph {pi}: \"{text[:80]}\"")
                log(f"  Style: {para.style.name if para.style else 'None'}")
                
                # Duyệt từng run
                for ri, run in enumerate(para.runs):
                    run_text = run.text
                    if not run_text:
                        continue
                    
                    # Formatting qua python-docx API
                    log(f"\n    Run {ri}: \"{run_text}\"")
                    
                    # Color
                    try:
                        color_rgb = run.font.color.rgb
                        color_type = run.font.color.type
                        color_theme = run.font.color.theme_color
                        log(f"      color.rgb = {color_rgb}")
                        log(f"      color.type = {color_type}")
                        log(f"      color.theme_color = {color_theme}")
                    except Exception as e:
                        log(f"      color: ERROR - {e}")
                    
                    # Underline
                    try:
                        underline = run.font.underline
                        log(f"      underline = {underline}")
                    except Exception as e:
                        log(f"      underline: ERROR - {e}")
                    
                    # Bold
                    try:
                        bold = run.font.bold
                        log(f"      bold = {bold}")
                    except Exception as e:
                        log(f"      bold: ERROR - {e}")
                    
                    # Raw XML
                    rpr = run._element.find(qn('w:rPr'))
                    if rpr is not None:
                        xml_str = etree.tostring(rpr, pretty_print=True).decode('utf-8')
                        log(f"      Raw XML rPr:")
                        for line in xml_str.strip().split('\n'):
                            log(f"        {line}")
                    else:
                        log(f"      Raw XML rPr: None")
                    
                    # Kiểm tra style inheritance
                    style_font = None
                    try:
                        if run.style and run.style.font:
                            style_font = run.style.font
                            log(f"      Style font color: {style_font.color.rgb if style_font.color else 'None'}")
                    except Exception:
                        pass
        
        # Dump paragraph style formatting
        log(f"\n\n{'='*70}")
        log(f"  PARAGRAPH STYLES USED:")
        log(f"{'='*70}")
        
        styles_seen = set()
        for para in doc.paragraphs:
            if para.style and para.style.name not in styles_seen:
                styles_seen.add(para.style.name)
                log(f"  - {para.style.name}")
                try:
                    sf = para.style.font
                    if sf.color and sf.color.rgb:
                        log(f"    font.color.rgb = {sf.color.rgb}")
                    if sf.color and sf.color.theme_color:
                        log(f"    font.color.theme_color = {sf.color.theme_color}")
                    if sf.underline:
                        log(f"    font.underline = {sf.underline}")
                except Exception:
                    pass
        
        # Check document theme colors
        log(f"\n\n{'='*70}")
        log(f"  DOCUMENT THEME/DEFAULT FORMAT:")
        log(f"{'='*70}")
        
        # Check default paragraph format
        try:
            body = doc.element.body
            # Find sectPr for page-level settings
            for child in body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'sectPr':
                    log(f"  sectPr found")
            
            # (Bạn có thể thêm logic phân tích styles_part tại đây nếu cần)
        except Exception as e:
            log(f"  (Lỗi khi đọc theme: {e})")

    finally:
        # Đảm bảo file được đóng an toàn sau khi chạy xong hoặc nếu có lỗi xảy ra
        if f_out:
            f_out.close()
            print(f"\n[INFO] Đã ghi toàn bộ log ra file: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        # Mặc định dùng file trong thư mục hiện tại
        filepath = os.path.join(os.path.dirname(__file__), "Các quy tắc trộn đề.docx")
        if not os.path.exists(filepath):
            print("Cách dùng: python debug_formatting.py <file.docx>")
            sys.exit(1)
    else:
        filepath = sys.argv[1]
    
    if not os.path.exists(filepath):
        print(f"❌ File không tồn tại: {filepath}")
        sys.exit(1)
    
    dump_paragraph_formatting(filepath, output_path="Test.txt")
    
    print("\n✅ Debug hoàn tất")
