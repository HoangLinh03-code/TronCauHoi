# -*- coding: utf-8 -*-
"""
Parser — Parse file DOCX đề thi, trích xuất câu hỏi + đáp án + nhận diện đáp án đúng.

Chạy debug:
    python parser.py input/de_goc.docx
"""

import os
import sys
import re
from copy import deepcopy
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import RGBColor
from lxml import etree

from models import ExamDocument, Question, OptionData, QuestionBlock

# Đảm bảo encoding UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')


# ============================================================
# Constants & Patterns
# ============================================================

# Namespace XML
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

# Pattern nhận diện câu hỏi
QUESTION_PATTERNS = [
    re.compile(r'(?:Câu|Question)\s*(\d+)\s*[.:)]\s*', re.IGNORECASE),
    re.compile(r'^(\d+)\s*[.:)]\s+', re.MULTILINE),
]

# Pattern nhận diện đáp án (mỗi đáp án 1 dòng riêng)
OPTION_PATTERN = re.compile(r'^([A-D])\s*[.)]\s*(.*)', re.IGNORECASE)
OPTION_INLINE_PATTERN = re.compile(r'([A-D])\s*[.)]\s*(.*?)(?=\s+[A-D]\s*[.)]\s|$)', re.IGNORECASE)

# Pattern nhận diện option letter trong run text: "A. ", "B.", "C) ", etc.
OPTION_LETTER_RUN_PATTERN = re.compile(r'^([A-D])\s*[.)]\s*$', re.IGNORECASE)

# Pattern nhận diện section headers
SECTION_PATTERNS = [
    re.compile(r'(?i)^(?:PART|SECTION|PHẦN)\s*(\d+|[IVXLC]+)\s*[.:]*\s*(.*)', re.IGNORECASE),
    re.compile(r'^(?:I{1,3}|IV|V|VI{0,3})\s*[.:]\s*(.*)', re.IGNORECASE),
]

# Pattern nhận diện context range (đoạn văn đọc hiểu)
CONTEXT_RANGE_PATTERN = re.compile(
    r'(?:questions?|câu)\s*(\d+)\s*[-–to]+\s*(\d+)',
    re.IGNORECASE
)


# ============================================================
# XML Helpers
# ============================================================

def get_paragraph_text(para_element) -> str:
    """Lấy text thuần từ 1 paragraph element (bao gồm text trong math)."""
    parts = []
    for child in para_element.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 't' and child.text:
            parts.append(child.text)
    return ''.join(parts).strip()


def get_run_formatting(run_element) -> dict:
    """Lấy thông tin formatting của 1 run element."""
    info = {
        'text': '',
        'is_red': False,
        'is_underline': False,
        'is_bold': False,
    }

    # Lấy text
    for t in run_element.iter():
        tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
        if tag == 't' and t.text:
            info['text'] += t.text

    # Lấy run properties
    rpr = run_element.find(f'{{{WML_NS}}}rPr')
    if rpr is not None:
        # Kiểm tra color
        color_el = rpr.find(f'{{{WML_NS}}}color')
        if color_el is not None:
            color_val = color_el.get(f'{{{WML_NS}}}val', '')
            if not color_val:
                color_val = color_el.get('val', '')  # Fallback không namespace
            if color_val:
                info['is_red'] = _is_red_color(color_val)

        # Kiểm tra underline
        u_el = rpr.find(f'{{{WML_NS}}}u')
        if u_el is not None:
            u_val = u_el.get(f'{{{WML_NS}}}val', '')
            if not u_val:
                u_val = u_el.get('val', '')
            if u_val and u_val.lower() != 'none':
                info['is_underline'] = True

        # Kiểm tra bold
        b_el = rpr.find(f'{{{WML_NS}}}b')
        if b_el is not None:
            b_val = b_el.get(f'{{{WML_NS}}}val', '')
            if not b_val:
                b_val = b_el.get('val', '')
            # <w:b/> hoặc <w:b val="true"> → bold
            if b_val.lower() not in ('false', '0'):
                info['is_bold'] = True

    return info


def _is_red_color(hex_color: str) -> bool:
    """Kiểm tra 1 mã hex có phải màu đỏ không."""
    hex_color = hex_color.strip().lstrip('#')
    if len(hex_color) != 6:
        return False
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return r > 180 and g < 100 and b < 100
    except ValueError:
        return False


def _detect_correct_by_option_letter_run(para_element) -> Optional[str]:
    """
    CHIẾN LƯỢC CHÍNH: Tìm đáp án đúng bằng cách kiểm tra run chứa 
    chữ cái đáp án (VD: "A. ", "B. ", "C. ", "D. ").
    
    Trong đề thi, run chứa chữ cái đáp án đúng sẽ có font.color = đỏ (FF0000).
    Underline KHÔNG phải dấu hiệu đáp án đúng (vì dùng cho phần phát âm).
    
    Returns: 'A'/'B'/'C'/'D' hoặc None
    """
    correct_letter = None
    red_letters = []  # Có thể có nhiều run, nhưng chỉ 1 letter nên cần gom
    
    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        
        fmt = get_run_formatting(child)
        text = fmt['text'].strip()
        
        if not text:
            continue
        
        # Kiểm tra xem run này có phải chứa "A." / "B." / "C." / "D." không
        m = OPTION_LETTER_RUN_PATTERN.match(text)
        if m and fmt['is_red']:
            letter = m.group(1).upper()
            red_letters.append(letter)
        
        # Trường hợp 2: Run chứa nội dung đáp án (không chỉ letter) cũng đỏ
        # VD: run text = "become " với color = red
        # → Nó thuộc về option nào? Cần biết option nào đang active
        
    if red_letters:
        # Nếu có nhiều letter đỏ → lấy cái đầu tiên (bất thường)
        correct_letter = red_letters[0]
    
    return correct_letter


def _detect_correct_by_content_runs(para_element) -> Optional[str]:
    """
    CHIẾN LƯỢC PHỤ: Nếu không phát hiện qua option letter run,
    thì duyệt tất cả runs, theo dõi "đang ở option nào" dựa vào 
    tab/chữ cái tiêu đề, rồi kiểm tra run nội dung có đỏ không.
    
    Logic:
    - Khi gặp run chứa "A. " → bắt đầu option A
    - Khi gặp run chứa "B. " → bắt đầu option B  
    - ... 
    - Nếu run nội dung (không phải letter) có color đỏ → option đó là đáp án đúng
    
    Returns: 'A'/'B'/'C'/'D' hoặc None
    """
    current_option = None
    red_content_options = {}  # {letter: red_char_count}
    
    letter_pattern = re.compile(r'([A-D])\s*[.)]\s*', re.IGNORECASE)
    
    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        
        fmt = get_run_formatting(child)
        text = fmt['text']
        
        if not text:
            continue
        
        # Kiểm tra xem run có chứa option letter không
        m_letter = letter_pattern.search(text)
        if m_letter:
            letter = m_letter.group(1).upper()
            current_option = letter
            
            # Nếu chính run chứa letter cũng đỏ → đáp án đúng
            if fmt['is_red']:
                red_content_options[letter] = red_content_options.get(letter, 0) + len(text.strip())
            continue
        
        # Tab hoặc whitespace → không đổi option, bỏ qua
        if text.strip() == '' or text.strip() == '\t':
            continue
        
        # Run nội dung thuộc option hiện tại
        if current_option and fmt['is_red']:
            red_content_options[current_option] = red_content_options.get(current_option, 0) + len(text.strip())
    
    if red_content_options:
        # Chỉ đáp án đúng khi CHỈ 1 option có đỏ
        if len(red_content_options) == 1:
            return list(red_content_options.keys())[0]
        # Nếu nhiều option đỏ → lấy option có NHIỀU ký tự đỏ nhất
        # (vì đáp án đúng thường toàn bộ đỏ, trong khi lỗi ngẫu nhiên chỉ 1-2 char)
        max_letter = max(red_content_options, key=red_content_options.get)
        return max_letter
    
    return None


def check_paragraph_correct_separate(para_element) -> Tuple[bool, bool]:
    """
    Kiểm tra formatting cho trường hợp đáp án trên DÒNG RIÊNG.
    Trả về (has_red, has_underline_only).
    
    Đối với đáp án trên dòng riêng (VD paragraph text = "A. became"):
    - Red → chắc chắn là đáp án đúng
    - Underline → CÓ THỂ là đáp án đúng (cần verify)
    """
    has_red = False
    has_underline_only = False
    
    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            fmt = get_run_formatting(child)
            if fmt['text'].strip():
                if fmt['is_red']:
                    has_red = True
                if fmt['is_underline'] and not fmt['is_red']:
                    has_underline_only = True
    
    return has_red, has_underline_only


def check_paragraph_correct_docx(paragraph) -> Tuple[bool, bool]:
    """Kiểm tra đáp án đúng qua python-docx API. Returns (has_red, has_underline)."""
    has_red = False
    has_underline = False
    
    for run in paragraph.runs:
        if not run.text.strip():
            continue

        # Kiểm tra màu đỏ
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                r, g, b = rgb[0], rgb[1], rgb[2]
                if r > 180 and g < 100 and b < 100:
                    has_red = True
        except Exception:
            pass

        # Kiểm tra gạch chân
        try:
            if run.font.underline is not None and run.font.underline is not False:
                has_underline = True
        except Exception:
            pass

    return has_red, has_underline


# ============================================================
# Main Parser
# ============================================================

class DocxParser:
    """Parse file DOCX đề thi → ExamDocument."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.exam = ExamDocument(filepath=filepath)
        self._paragraphs = []  # (paragraph_obj, element, text)

    def parse(self) -> ExamDocument:
        """Parse toàn bộ file DOCX."""
        print(f"\n📄 Đang parse: {os.path.basename(self.filepath)}")

        # Bước 1: Thu thập tất cả paragraphs
        self._collect_paragraphs()
        print(f"   Tổng paragraph: {len(self._paragraphs)}")

        # Bước 2: Tìm vị trí bắt đầu câu hỏi
        question_positions = self._find_question_positions()
        print(f"   Vị trí câu hỏi phát hiện: {len(question_positions)}")

        if not question_positions:
            print("   ⚠️  Không tìm thấy câu hỏi nào!")
            return self.exam

        # Bước 3: Tách header (trước câu hỏi đầu tiên)
        first_q_idx = question_positions[0][0]
        for i in range(first_q_idx):
            para_obj, elem, text = self._paragraphs[i]
            self.exam.header_elements.append(elem)
            self.exam.header_text += text + "\n"

        # Bước 4: Parse từng câu hỏi
        for qi in range(len(question_positions)):
            q_idx, q_num = question_positions[qi]

            # Xác định phạm vi paragraph của câu này
            if qi + 1 < len(question_positions):
                next_q_idx = question_positions[qi + 1][0]
            else:
                next_q_idx = len(self._paragraphs)

            q_paragraphs = self._paragraphs[q_idx:next_q_idx]
            question = self._parse_single_question(q_num, q_paragraphs)

            if question:
                self.exam.questions.append(question)

        self.exam.total_questions = len(self.exam.questions)

        # Bước 5: Detect context blocks (đoạn văn đọc hiểu)
        self._detect_context_blocks()

        print(f"   ✅ Parse xong: {self.exam.total_questions} câu hỏi")
        return self.exam

    def _collect_paragraphs(self):
        """Thu thập tất cả paragraph + element + text."""
        for i, para in enumerate(self.doc.paragraphs):
            elem = para._element
            text = para.text.strip()
            self._paragraphs.append((para, elem, text))

    def _find_question_positions(self) -> List[Tuple[int, int]]:
        """Tìm vị trí (index, số câu) của tất cả câu hỏi."""
        positions = []
        for i, (para, elem, text) in enumerate(self._paragraphs):
            if not text:
                continue
            for pattern in QUESTION_PATTERNS:
                m = pattern.match(text)
                if m:
                    q_num = int(m.group(1))
                    # Tránh trùng: chỉ lấy nếu số câu chưa có hoặc tăng dần
                    if not positions or q_num > positions[-1][1]:
                        positions.append((i, q_num))
                    elif q_num == positions[-1][1]:
                        pass  # Bỏ qua duplicate
                    break
        return positions

    def _parse_single_question(self, q_num: int,
                                q_paragraphs: list) -> Optional[Question]:
        """Parse 1 câu hỏi từ danh sách paragraphs."""
        question = Question(original_number=q_num)

        # Tách stem và options
        stem_paras = []
        option_paras = {}  # {'A': [(para, elem, text)], ...}
        current_option = None
        context_paras = []  # Paragraphs trước stem (context/đọc hiểu)

        # Kiểm tra xem đáp án nằm trên dòng riêng hay cùng dòng
        all_texts = [t for _, _, t in q_paragraphs]
        inline_options = self._detect_inline_options(all_texts)

        if inline_options:
            # Đáp án nằm trên cùng dòng hoặc 2 dòng
            return self._parse_question_inline_options(q_num, q_paragraphs, inline_options)

        # Trường hợp thông thường: mỗi đáp án 1 dòng
        found_stem = False
        for para, elem, text in q_paragraphs:
            if not text:
                continue

            # Kiểm tra xem có phải đầu đáp án không
            opt_match = OPTION_PATTERN.match(text)
            if opt_match:
                letter = opt_match.group(1).upper()
                current_option = letter
                if letter not in option_paras:
                    option_paras[letter] = []
                option_paras[letter].append((para, elem, text))
                found_stem = True  # Đã qua phần stem
            elif current_option:
                # Continuation của đáp án trước
                option_paras[current_option].append((para, elem, text))
            else:
                # Phần stem câu hỏi
                stem_paras.append((para, elem, text))
                found_stem = True

        # Build Question object
        if stem_paras:
            question.stem_text = ' '.join(t for _, _, t in stem_paras)
            question.stem_elements = [elem for _, elem, _ in stem_paras]

            # Lấy text thực của câu hỏi (bỏ prefix "Câu X:")
            for pattern in QUESTION_PATTERNS:
                question.stem_text = pattern.sub('', question.stem_text).strip()

        # Build options — Cho trường hợp DÒNG RIÊNG
        for letter in ['A', 'B', 'C', 'D']:
            if letter in option_paras:
                paras = option_paras[letter]
                opt_text = ' '.join(t for _, _, t in paras)
                # Bỏ prefix "A. " / "A) "
                opt_text = re.sub(r'^[A-D]\s*[.)]\s*', '', opt_text).strip()

                # Kiểm tra formatting đáp án đúng — dòng riêng
                # Ưu tiên: RED > underline
                is_correct = False
                for para_obj, elem, _ in paras:
                    has_red, has_underline = check_paragraph_correct_separate(elem)
                    if has_red:
                        is_correct = True
                        break
                
                # Nếu không có red, thử underline (CHỈ khi đáp án trên dòng riêng)
                if not is_correct:
                    for para_obj, elem, _ in paras:
                        has_red, has_underline = check_paragraph_correct_separate(elem)
                        if has_underline:
                            is_correct = True
                            break
                
                # Backup: python-docx API
                if not is_correct:
                    for para_obj, _, _ in paras:
                        has_red, has_underline = check_paragraph_correct_docx(para_obj)
                        if has_red:
                            is_correct = True
                            break

                question.options[letter] = OptionData(
                    original_letter=letter,
                    current_letter=letter,
                    text=opt_text,
                    is_correct=is_correct,
                    elements=[elem for _, elem, _ in paras],
                )

        # Nếu dùng underline → verify chỉ 1 đáp án đúng
        correct_letters = [l for l in 'ABCD' if l in question.options and question.options[l].is_correct]
        if len(correct_letters) > 1:
            # Nếu nhiều đáp án đúng → chỉ giữ những đáp án có RED
            red_letters = []
            for letter in correct_letters:
                paras = option_paras.get(letter, [])
                for _, elem, _ in paras:
                    has_red, _ = check_paragraph_correct_separate(elem)
                    if has_red:
                        red_letters.append(letter)
                        break
            
            if red_letters:
                # Reset tất cả → chỉ giữ red
                for letter in 'ABCD':
                    if letter in question.options:
                        question.options[letter].is_correct = (letter in red_letters)
            else:
                # Không có red → giữ tất cả underline (có thể là dạng đề phát âm)
                # → reset tất cả, để user fix thủ công
                for letter in 'ABCD':
                    if letter in question.options:
                        question.options[letter].is_correct = False

        # Xác định đáp án đúng
        correct = question.get_correct_letter()
        if correct:
            question.correct_answer = correct
            question.original_correct_answer = correct

        return question

    def _detect_inline_options(self, texts: List[str]) -> Optional[dict]:
        """
        Phát hiện đáp án nằm trên cùng 1 dòng (hoặc 2 dòng).
        Trả về dict {line_index: [(letter, text), ...]} hoặc None.
        """
        for i, text in enumerate(texts):
            matches = list(OPTION_INLINE_PATTERN.finditer(text))
            if len(matches) >= 3:
                options = [(m.group(1).upper(), m.group(2).strip()) for m in matches]
                return {'line_index': i, 'options': options, 'type': 'single_line'}

        # Kiểm tra 2 dòng liên tiếp, mỗi dòng 2 đáp án
        for i in range(len(texts) - 1):
            matches1 = list(OPTION_INLINE_PATTERN.finditer(texts[i]))
            matches2 = list(OPTION_INLINE_PATTERN.finditer(texts[i + 1]))
            if len(matches1) == 2 and len(matches2) == 2:
                options = (
                    [(m.group(1).upper(), m.group(2).strip()) for m in matches1] +
                    [(m.group(1).upper(), m.group(2).strip()) for m in matches2]
                )
                return {'line_index': i, 'options': options, 'type': 'two_lines'}

        return None

    def _parse_question_inline_options(self, q_num: int,
                                        q_paragraphs: list,
                                        inline_info: dict) -> Question:
        """
        Parse câu hỏi khi đáp án nằm inline (cùng dòng).
        
        CHIẾN LƯỢC MỚI: Duyệt runs trong paragraph, tìm run chứa 
        "A. " / "B. " / "C. " / "D. " → kiểm tra run đó có đỏ không.
        KHÔNG dùng underline cho inline options (vì underline = phần phát âm).
        """
        question = Question(original_number=q_num, is_inline=True)
        line_idx = inline_info['line_index']
        options_data = inline_info['options']

        # Stem = tất cả paragraphs trước dòng chứa đáp án
        stem_paras = q_paragraphs[:line_idx]
        if not stem_paras and line_idx < len(q_paragraphs):
            # Đáp án trên cùng dòng với câu hỏi → tách phần stem
            para, elem, text = q_paragraphs[line_idx]
            first_opt_match = OPTION_INLINE_PATTERN.search(text)
            if first_opt_match:
                stem_text = text[:first_opt_match.start()].strip()
                question.stem_text = stem_text
                question.stem_elements = [elem]
            stem_paras = []
        else:
            question.stem_text = ' '.join(t for _, _, t in stem_paras)
            question.stem_elements = [elem for _, elem, _ in stem_paras]

        # Bỏ prefix "Câu X:"
        for pattern in QUESTION_PATTERNS:
            question.stem_text = pattern.sub('', question.stem_text).strip()

        # === TÌM ĐÁP ÁN ĐÚNG bằng cách duyệt RUNS ===
        opt_line_idx = line_idx
        opt_end_idx = line_idx + 1 if inline_info['type'] == 'single_line' else line_idx + 2
        
        # Chiến lược 1: Tìm run chứa "A. " có đỏ
        correct_from_letter_run = None
        for pi in range(opt_line_idx, min(opt_end_idx, len(q_paragraphs))):
            _, elem, _ = q_paragraphs[pi]
            found = _detect_correct_by_option_letter_run(elem)
            if found:
                correct_from_letter_run = found
                break
        
        # Chiến lược 2: Duyệt content runs có đỏ
        correct_from_content = None
        if not correct_from_letter_run:
            for pi in range(opt_line_idx, min(opt_end_idx, len(q_paragraphs))):
                _, elem, _ = q_paragraphs[pi]
                found = _detect_correct_by_content_runs(elem)
                if found:
                    correct_from_content = found
                    break
        
        correct_letter = correct_from_letter_run or correct_from_content

        # Build options
        opt_elements = [
            elem for _, elem, _ in q_paragraphs[opt_line_idx:opt_end_idx]
        ]

        for letter, opt_text in options_data:
            is_correct = (letter == correct_letter) if correct_letter else False

            question.options[letter] = OptionData(
                original_letter=letter,
                current_letter=letter,
                text=opt_text,
                is_correct=is_correct,
                elements=opt_elements,
            )

        if correct_letter:
            question.correct_answer = correct_letter
            question.original_correct_answer = correct_letter

        return question

    def _detect_context_blocks(self):
        """Phát hiện context blocks (đoạn văn đọc hiểu gắn liền câu hỏi)."""
        block_id = 0

        # Tìm trong header và giữa các câu hỏi
        all_text = self.exam.header_text
        for q in self.exam.questions:
            if q.context_text:
                all_text += q.context_text

        # Tìm pattern "questions X-Y" hoặc "câu X đến Y"
        for q in self.exam.questions:
            for elem in q.stem_elements:
                text = get_paragraph_text(elem)
                m = CONTEXT_RANGE_PATTERN.search(text)
                if m:
                    start_q = int(m.group(1))
                    end_q = int(m.group(2))
                    block = QuestionBlock(
                        block_id=block_id,
                        context_text=text,
                        question_numbers=list(range(start_q, end_q + 1))
                    )
                    self.exam.blocks.append(block)
                    # Gán block_id cho các câu thuộc block
                    for qn in range(start_q, end_q + 1):
                        qq = self.exam.get_question(qn)
                        if qq:
                            qq.block_id = block_id
                    block_id += 1

    def detect_sections(self) -> List[dict]:
        """Phát hiện section headers trong đề (PART 1, Section A, ...)."""
        sections = []
        for i, (para, elem, text) in enumerate(self._paragraphs):
            if not text:
                continue
            for pattern in SECTION_PATTERNS:
                m = pattern.match(text)
                if m:
                    sections.append({
                        'index': i,
                        'text': text,
                        'match': m.group(0),
                    })
                    break
        return sections


def parse_docx(filepath: str) -> ExamDocument:
    """Convenience function: Parse 1 file DOCX."""
    parser = DocxParser(filepath)
    return parser.parse()


# ============================================================
# Debug: chạy python parser.py <file.docx>
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("  DEBUG: parser.py — Parse DOCX đề thi")
    print("=" * 60)

    if len(sys.argv) < 2:
        print("\nCách dùng: python parser.py <file.docx>")
        print("Ví dụ:     python parser.py input/de_goc.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"\n❌ File không tồn tại: {filepath}")
        sys.exit(1)

    # Parse
    exam = parse_docx(filepath)

    # Hiển thị kết quả
    print("\n" + "=" * 60)
    print(exam.summary())
    print("=" * 60)

    # Hiển thị từng câu hỏi
    for q in exam.questions:
        print(f"\n--- Câu {q.original_number} ---")
        print(f"  Stem: {q.stem_text[:100]}")
        if q.context_text:
            print(f"  Context: {q.context_text[:80]}...")
        if q.block_id is not None:
            print(f"  Block ID: {q.block_id}")

        for letter in 'ABCD':
            opt = q.options.get(letter)
            if opt:
                mark = "✓" if opt.is_correct else " "
                print(f"  [{mark}] {letter}. {opt.text[:80]}")
            else:
                print(f"  [ ] {letter}. (không tìm thấy)")

        if q.correct_answer:
            print(f"  → Đáp án đúng: {q.correct_answer}")
        else:
            print(f"  ⚠️  Không phát hiện đáp án đúng!")

    # Detect sections
    parser = DocxParser(filepath)
    parser._collect_paragraphs()
    sections = parser.detect_sections()
    if sections:
        print(f"\n--- Sections phát hiện ---")
        for s in sections:
            print(f"  [{s['index']}] {s['text']}")
    else:
        print(f"\n--- Không phát hiện section headers ---")

    # Thống kê
    correct_count = sum(1 for q in exam.questions if q.correct_answer)
    no_correct = [q.original_number for q in exam.questions if not q.correct_answer]
    print(f"\n📊 Thống kê:")
    print(f"   Tổng câu: {exam.total_questions}")
    print(f"   Có đáp án đúng: {correct_count}")
    if no_correct:
        print(f"   ⚠️  Không phát hiện đáp án: câu {no_correct}")
    if exam.blocks:
        print(f"   Blocks đọc hiểu: {len(exam.blocks)}")
        for b in exam.blocks:
            print(f"     Block {b.block_id}: câu {b.question_numbers}")

    print("\n✅ parser.py OK")
